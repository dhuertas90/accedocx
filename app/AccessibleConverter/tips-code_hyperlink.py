from docx import Document
import docx
from docx.opc.constants import RELATIONSHIP_TYPE as RT

def detectar_hipervinculos(documento):
    doc = Document(documento)

    dic_ids = {}
    for paragraph in doc.paragraphs:
        # Test hyperlink function
        # Iterator on Hyperlinks
        for relId, rel in doc._part.rels.items():
            # print(f"{relId} : {rel}")
            if rel.reltype == RT.HYPERLINK:
                if relId not in dic_ids.keys():
                   dic_ids[relId] = rel._target 
                # print(relId +" --> " + rel._target)
                # add_hyperlink(paragraph, relId, rel._target)
        print(dic_ids)
# Uso del ejemplo:
detectar_hipervinculos('origen.docx')

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink



# document = docx.Document()
# p = document.add_paragraph()

# #add a hyperlink with the normal formatting (blue underline)
# hyperlink = add_hyperlink(p, 'http://www.google.com', 'Google', None, True)

# #add a hyperlink with a custom color and no underline
# hyperlink = add_hyperlink(p, 'http://www.google.com', 'Google', 'FF8822', False)

# document.save('demo.docx')


# Busco hyperlink
# if bool(paragraph._element.xpath(".//w:hyperlink")):
#     print("Obtenemos un hyper: ")
#     print(paragraph._element.xpath(".//w:hyperlink"))
#     text = paragraph.text + "\n" + str(dic_hyper)


# # ejemplo
    # for para in doc.paragraphs:
    #     for link in para._element.xpath(".//w:hyperlink"):
    #         inner_run = link.xpath("w:r", namespaces=link.nsmap)[0]
    #         # print link text
    #         print(inner_run.text)
    #         # print link relationship id
    #         print(rId)
    #         # print link URL
    #         print(doc._part.rels[rId]._target)


      # Listado. Iterar sobre los items existentes, obtener los vinculos.
    # list_url_hyper = []
    # dic_hyper = {}
    # for relId, rel in doc._part.rels.items():
    #     if rel.reltype == RT.HYPERLINK:
    #         dic_hyper[relId] = rel._target
    #         # list_url_hyper.append({relId:rel._target})
    #         print("un hiper!")
    # print(dic_hyper)
    # Diccionario. Guardar 'rId:filenames' relaciones.